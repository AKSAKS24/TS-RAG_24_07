import os
from docx import Document
from docx.shared import Pt, RGBColor

def add_heading(doc, text: str, level: int):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    if level == 1:
        run.bold = True
        run.underline = True
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.font.size = Pt(14)
    elif level == 2:
        run.bold = True
        run.font.size = Pt(12)
    else:
        run.bold = True
        run.font.size = Pt(11)

def add_bullet_point(doc, text: str):
    paragraph = doc.add_paragraph(style='List Bullet')
    run = paragraph.add_run(text)
    run.font.size = Pt(11)

def add_paragraph(doc, text: str):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(11)

def add_markdown_table(doc, lines):
    if len(lines) < 3:
        return  # Not enough for markdown table
    headers = [cell.strip(" *") for cell in lines[0].split("|") if cell.strip()]
    rows = [
        [cell.strip() for cell in row.split("|") if cell.strip()]
        for row in lines[2:]
    ]
    col_count = len(headers)
    filtered_rows = [row for row in rows if len(row) == col_count]
    table = doc.add_table(rows=1, cols=col_count)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row in filtered_rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = cell

def write_docx_report(text: str, template: str, output_path: str) -> str:
    directory = os.path.dirname(output_path)
    if directory and not os.path.exists(directory):
        os.makedirs(directory)
    doc = Document()
    lines = text.splitlines()
    in_table = False
    table_lines = []
    markdown_table_pattern = lambda line: line.strip().startswith('|') and line.strip().endswith('|')
    for line in lines:
        stripped = line.strip()
        # Flush table if blank line or next non-table starts
        if not stripped:
            if in_table and table_lines:
                add_markdown_table(doc, table_lines)
                table_lines.clear()
                in_table = False
            continue
        if markdown_table_pattern(stripped):
            table_lines.append(stripped)
            in_table = True
            continue
        elif in_table and table_lines:
            add_markdown_table(doc, table_lines)
            table_lines.clear()
            in_table = False
        if stripped.startswith('###'):
            add_heading(doc, stripped.lstrip('#').strip(), level=3)
        elif stripped.startswith('##'):
            add_heading(doc, stripped.lstrip('#').strip(), level=2)
        elif stripped.startswith('#'):
            add_heading(doc, stripped.lstrip('#').strip(), level=1)
        elif stripped.startswith('-'):
            add_bullet_point(doc, stripped.lstrip('-').strip())
        else:
            add_paragraph(doc, stripped)
    if in_table and table_lines:
        add_markdown_table(doc, table_lines)
    doc.save(output_path)
    return output_path